#!/usr/bin/env python3
"""Generate split BRD documents from the standard Workfront-to-CTT integration document."""

from copy import deepcopy
from datetime import date

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SOURCE_PATH = "/workspace/20260729-Workfront to CTT Integration.docx"
BRD1_PATH = "/workspace/BRD-01 - CTT and CTT Request Portal FY27 Updates.docx"
BRD2_PATH = "/workspace/BRD-02 - Workfront to CTT Integration Feasibility Study.docx"


def set_cell_text(cell, text):
    cell.text = text


def add_para(doc, text, style="Cisco Normal"):
    return doc.add_paragraph(text, style=style)


def add_heading(doc, text, level="main"):
    style = {
        "section": "Cisco Heading",
        "main": "Cisco Main Heading",
        "sub": "Cisco Sub Heading",
    }[level]
    return doc.add_paragraph(text, style=style)


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="Cisco Bullet Point")


def add_list(doc, text):
    return doc.add_paragraph(text, style="Cisco List Paragraph")


def add_story(doc, story_id, text):
    p = doc.add_paragraph(style="Cisco Normal")
    run = p.add_run(f"{story_id}: ")
    run.bold = True
    p.add_run(text)


def copy_table(source_doc, table_index, target_doc):
    src = source_doc.tables[table_index]
    rows = len(src.rows)
    cols = len(src.columns)
    tbl = target_doc.add_table(rows=rows, cols=cols)
    tbl.style = src.style
    for r in range(rows):
        for c in range(cols):
            tbl.rows[r].cells[c].text = src.rows[r].cells[c].text
    return tbl


def build_revision_table(doc, doc_id, title, version, description):
    add_heading(doc, "Revision History", "section")
    table = doc.add_table(rows=2, cols=4)
    headers = ["Date", "Version", "Description", "Author"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h)
    set_cell_text(table.rows[1].cells[0], date.today().strftime("%d/%m/%Y"))
    set_cell_text(table.rows[1].cells[1], version)
    set_cell_text(table.rows[1].cells[2], description)
    set_cell_text(table.rows[1].cells[3], "Amanda Pattenden")
    doc.add_paragraph()


def build_metadata_block(doc, fields):
    table = doc.add_table(rows=len(fields), cols=2)
    for i, (k, v) in enumerate(fields):
        set_cell_text(table.rows[i].cells[0], k)
        set_cell_text(table.rows[i].cells[1], v)
    doc.add_paragraph()


def blank_document_from_template(template_path):
    """Create a new document that retains Cisco styles from the source template."""
    doc = Document(template_path)
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
            body.remove(child)
    return doc


def build_brd1(source):
    doc = blank_document_from_template(SOURCE_PATH)
    add_heading(doc, "Business Requirements Document", "section")
    build_metadata_block(
        doc,
        [
            ("Document Title", "CTT Tool and CTT Request Portal FY27 Updates"),
            ("Document ID", "BRD-WF-CTT-001"),
            ("Version", "1.0 (Draft)"),
            ("Status", "Draft — For Review"),
            ("Target Delivery", "Q1 FY27"),
            ("Author", "Amanda Pattenden"),
            ("Date", "July 2026"),
            (
                "Related Documents",
                "20260729-Workfront to CTT Integration.docx; BRD-02 - Workfront to CTT Integration Feasibility Study",
            ),
        ],
    )
    build_revision_table(
        doc,
        "BRD-WF-CTT-001",
        "CTT Tool and CTT Request Portal FY27 Updates",
        "1.0",
        "Initial BRD split from combined integration document — Q1 delivery track",
    )

    add_heading(doc, "Introduction", "section")
    add_heading(doc, "Executive Summary", "main")
    add_para(
        doc,
        "This Business Requirements Document defines the changes required to the Campaign Tagging and Tracking (CTT) tool, the CTT Request Portal, and OMS table publishing to support FY27 reporting requirements during the Workfront transition.",
    )
    add_para(
        doc,
        "Marketing has defined FY27 reporting data elements—including Business Priority, Stakeholder Unit, campaign hierarchy, funnel stage, and program context—that must be captured consistently across all marketing activities, including those that do not yet originate in Workfront.",
    )
    add_para(
        doc,
        "This BRD covers the minimum viable delivery required in Q1: updating CTT schemas, picklists, and request portal workflows, and publishing the updated attribute model to OMS tables so reporting and downstream consumers can access FY27-aligned data.",
    )
    add_para(
        doc,
        "Automated generation of Activity IDs (CCIDs) and Offer IDs (OIDs) from Workfront is explicitly out of scope for this document and is addressed separately in BRD-02 (targeted for Q2 at the earliest).",
    )

    add_heading(doc, "Problem Statement – Why?", "main")
    add_para(
        doc,
        "Adobe Workfront is the strategic platform for marketing work management, execution, and data governance. However, the CTT tool remains embedded in campaign identifier creation, metadata enrichment, routing, attribution, and reporting.",
    )
    add_para(
        doc,
        "Not all marketing activities currently flow through Workfront. Downstream reporting and attribution processes continue to rely on CTT-generated identifiers and metadata. Without immediate updates to CTT and the CTT Request Portal, FY27 reporting attributes will be captured inconsistently, creating reporting gaps and reducing data quality.",
    )
    add_para(
        doc,
        "Even if Workfront integration is deferred, the organization must still modernize the CTT ecosystem so marketing teams can manually align data across Workfront-managed and non-Workfront activities using the same FY27 attribute model.",
    )

    add_heading(doc, "Vision – What?", "main")
    add_para(
        doc,
        "Update the CTT tool and CTT Request Portal directly with all agreed FY27 fields, hierarchies, and picklists, and publish those values to OMS tables for centralized reporting consumption.",
    )
    add_bullet(doc, "Align Activity ID and Offer ID schemas to FY27 reporting standards.")
    add_bullet(doc, "Update CTT Request Portal Create New, Create Bulk, and Update Existing workflows.")
    add_bullet(doc, "Publish new and updated field values to OMS tables; deprecate legacy fields no longer required.")
    add_bullet(doc, "Enable manual data alignment for all marketing teams during the Workfront transition.")
    add_para(
        doc,
        "This solution is a temporary compatibility layer—not a long-term architecture. It bridges current-state dependencies while Workfront adoption expands and while automated integration is assessed in BRD-02.",
    )

    add_heading(doc, "Benefit Summary – Who?", "main")
    add_list(doc, "Marketing Teams: Request and manage IDs through updated portal workflows aligned to FY27 taxonomy.")
    add_list(doc, "Marketing Analytics & Insights Teams: Consistent FY27 attributes across Workfront and non-Workfront activity.")
    add_list(doc, "Reporting & Attribution Teams: OMS-published data supports centralized FY27 reporting validation.")
    add_list(doc, "CTT Administrators: Standardized schemas, picklists, and governance for bridge-period operations.")
    add_list(doc, "CRM / OCC Teams: No Salesforce tracking parameter changes required.")

    add_heading(doc, "Scope and Limitations", "section")
    add_heading(doc, "In-Scope", "main")
    add_bullet(doc, "Definition of FY27 reporting attributes and corresponding CTT field changes.")
    add_bullet(doc, "Activity ID schema updates: Stakeholder Unit rename, Primary Technology, Funnel Stage, Technology > Campaign > Program hierarchy.")
    add_bullet(doc, "Offer ID schema updates: Technology/Campaign/Program, Stakeholder Group, Funnel Stage, standardized Offer Type, updated Business Entity list.")
    add_bullet(doc, "Removal of legacy Vertical Market and Buying Centre fields from Offer ID and Activity ID workflows.")
    add_bullet(doc, "CTT Request Portal updates across Create New, Create Bulk, and Update Existing workflows.")
    add_bullet(doc, "Publication of all new and updated CTT field values to OMS tables.")
    add_bullet(doc, "Exclusion of new CTT field values from Tray.io, Allocadia, and SFDC API integrations during the bridge validation period.")

    add_heading(doc, "Out of Scope", "main")
    add_list(doc, "Automated Activity ID or Offer ID creation triggered from Workfront (see BRD-02).")
    add_list(doc, "Bidirectional synchronization between Workfront and CTT.")
    add_list(doc, "Full CTT decommissioning.")
    add_list(doc, "Salesforce, Eloqua, or OCC configuration changes.")
    add_list(doc, "Enterprise-wide Workfront adoption or process redesign.")
    add_list(doc, "Historical OCID data migration or enterprise-wide legacy ID cleanup.")

    add_heading(doc, "Assumptions", "main")
    add_bullet(doc, "FY27 reporting attributes are defined and sufficiently stable for configuration.")
    add_bullet(doc, "CTT can be configured to store and manage additional FY27 reporting attributes.")
    add_bullet(doc, "OMS ingestion pipelines can be updated to publish new fields and deprecate Vertical Market.")
    add_bullet(doc, "Delivery can proceed without Salesforce or OCC changes.")
    add_bullet(doc, "Non-Workfront teams will continue using the CTT Request Portal during transition.")

    add_heading(doc, "Dependencies", "main")
    add_bullet(doc, "Final approval of FY27 attributes, field mappings, picklists, and success criteria.")
    add_bullet(doc, "Timely deployment of schema updates and picklists in the core CTT database.")
    add_bullet(doc, "Deployment of updated form schemas and bulk templates in the CTT Request Portal.")
    add_bullet(doc, "OMS pipeline updates to ingest new CTT field values and process upserts.")
    add_bullet(doc, "Dedicated resources from CTT, Integration, Marketing Operations, and Reporting teams.")

    add_heading(doc, "Success Metrics", "main")
    add_list(doc, "Data Quality: ≥95% completeness for required FY27 attributes on new and updated CTT records.")
    add_list(doc, "OMS Availability: Updated CTT attribute model published to OMS before reporting validation.")
    add_list(doc, "Portal Alignment: CTT Request Portal workflows reflect updated FY27 taxonomy across all request types.")
    add_list(doc, "Reporting Consistency: Workfront and non-Workfront activities reportable using the same FY27 data elements.")
    add_list(doc, "Business Continuity: Existing CTT-dependent processes continue to function during schema changes.")

    add_heading(doc, "Stakeholder Requirements", "section")
    add_heading(doc, "Campaign Tagging & Tracking Tool", "main")
    add_heading(doc, "General Overview", "sub")
    add_story(
        doc,
        "CTT001",
        "As a CTT Administrator, I want CTT configured with interim bridge fields to carry Workfront-aligned FY27 data attributes, so that non-Workfront teams can use IDs with the same reporting data elements during transition (full list in Table 1 - ID Field List).",
    )
    add_story(
        doc,
        "CTT004",
        "As a CTT Administrator, I want existing CTT-dependent processes to continue functioning during the pilot, so that business continuity is protected while the bridge is validated.",
    )

    add_heading(doc, "Activity IDs - Additions", "sub")
    add_story(
        doc,
        "CTT005",
        "As a Marketing Manager, I want a Funnel Stage field added to Activity ID so that marketing activities are tagged to the correct funnel position for FY27 reporting and attribution (Table 4 - Funnel Stage).",
    )

    add_heading(doc, "Activity IDs – Amend", "sub")
    add_story(
        doc,
        "CTT006",
        "As a Data Governance Lead, I want the Activity ID field 'Business Unit' renamed to 'Stakeholder Unit' with a standardized list of values so that FY27 reporting uses consistent stakeholder attribution (Table 2 - Stakeholder List).",
    )
    add_story(
        doc,
        "CTT007",
        "As a Marketing Manager, I want to select a Primary Technology when creating an Activity ID so that only the relevant campaigns and associated programs are available for selection (Table 3 - Technology > Campaign > Program Mapping).",
    )

    add_heading(doc, "Activity IDs – Remove", "sub")
    add_story(
        doc,
        "CTT008",
        "As a Marketing Manager, I want the linkage between Campaign and Buying Centre removed from the Activity ID UI so that campaigns and programs can be selected from a distinct list without unnecessary buying centre constraints.",
    )

    add_heading(doc, "Offer IDs - Additions", "sub")
    add_story(
        doc,
        "CTT009",
        "As a Marketing Manager, I want new Technology, Campaign, and Program fields added to Offer ID, aligned to the same hierarchy and selection logic used on Activity ID (Table 3).",
    )
    add_story(
        doc,
        "CTT010",
        "As a Marketing Manager, I want a new Stakeholder Group field added to Offer ID with values aligned to the approved stakeholder hierarchy (Table 2).",
    )
    add_story(
        doc,
        "CTT011",
        "As a Marketing Manager, I want a new Funnel Stage field added to Offer ID with values aligned to the approved Funnel Stage framework (Table 4).",
    )

    add_heading(doc, "Offer IDs - Amends", "sub")
    add_story(
        doc,
        "CTT012",
        "As a Marketing Manager, I want the Offer Type field updated to a standardized list of values organized by Offer Category (Table 5 - Offer Category > Offer Type).",
    )
    add_story(
        doc,
        "CTT013",
        "As a CTT Administrator, I want the Business Entity field updated with the latest master list from Finance (Table 6 - Business and Sub Business Entity).",
    )

    add_heading(doc, "Offer IDs – Remove", "sub")
    add_story(
        doc,
        "CTT014",
        "As a Data Engineer, I want the Vertical Market field deprecated from the Offer ID ingestion pipeline to the OMS tables, so that legacy, unutilized data is not published to centralized reporting tables.",
    )

    add_heading(doc, "Data Capture", "sub")
    add_story(
        doc,
        "CTT015",
        "As a Data Engineer, I want all new and updated CTT field values published to OMS tables, so that reporting and downstream data consumers can access the updated attribute model.",
    )
    add_story(
        doc,
        "CTT016",
        "As an Integration Engineer, I want new CTT field values excluded from Tray.io, Allocadia, and SFDC API integrations during the bridge period, so that existing downstream integrations are not disrupted while CTT field changes are validated.",
    )

    add_heading(doc, "Campaign Tagging & Tracking Request Portal", "main")
    add_heading(doc, "Activity IDs", "sub")
    for story in [
        (
            "POR001",
            "As a Marketing Manager, I want the CTT Request Portal Stakeholder Group field updated to use the same standardized values as CTT Stakeholder Unit across Create New, Create Bulk, and Update Existing Activity ID request types (Table 2).",
        ),
        (
            "POR002",
            "As a Marketing Manager, I want the CTT Request Portal selection hierarchy changed from Buying Centre > Campaign Theme > Program to Technology > Campaign > Program across all Activity ID request types (Table 3).",
        ),
        (
            "POR003",
            "As a Marketing Manager, I want a Funnel Stage field added to Activity ID requests with values Awareness, Consideration, Evaluation, and Decision across all Activity ID request types (Table 4).",
        ),
    ]:
        add_story(doc, story[0], story[1])

    add_heading(doc, "Offer IDs", "sub")
    for story in [
        (
            "POR004",
            "As a Marketing Manager, I want the CTT Request Portal updated with the new standardized Offer Type list across all Offer ID request types (Table 5).",
        ),
        (
            "POR005",
            "As a Marketing Manager, I want the Business and Sub Business Entity list of values updated across all Offer ID request types (Table 6).",
        ),
        (
            "POR006",
            "As a Marketing Manager, I want the Vertical Market field removed from all Offer ID request types and templates.",
        ),
        (
            "POR007",
            "As a Marketing Manager, I want the Buying Centre field removed from all Offer ID request types and templates, and the logic to automatically populate the originating Activity ID removed, so that data from the originating Activity is captured directly within the Offer ID.",
        ),
        (
            "POR008",
            "As a Marketing Manager, I want the Technology > Campaign > Program hierarchy fields added to all Offer ID request types and templates (Table 3).",
        ),
        (
            "POR009",
            "As a Marketing Manager, I want the Funnel Stage field added to all Offer ID request types and templates (Table 4).",
        ),
    ]:
        add_story(doc, story[0], story[1])

    add_heading(doc, "Salesforce", "main")
    add_story(
        doc,
        "SF001",
        "As a CRM Product Owner, I want this initiative to require no immediate SFDC tracking parameter changes, so that One Cisco CRM work is not disrupted during the bridge period.",
    )
    add_story(
        doc,
        "SF002",
        "As a Data Engineer, I want existing SFDC campaign alignment to Activity ID and Offer ID to continue unchanged, so that lead creation and campaign member processes are not impacted.",
    )

    add_heading(doc, "Eloqua", "main")
    add_story(
        doc,
        "EL001",
        "As a Marketing Automation Specialist, I want no changes to Eloqua picklists, forms, or AOJ integrations during this phase, so that existing automation continues while CTT schema changes are validated.",
    )

    add_heading(doc, "CDF", "main")
    add_story(
        doc,
        "CDF001",
        "As a Reporting Analyst, I want this initiative to document how bridge-period data relates to current CDF Transaction Table dependencies, so that the decommission backlog includes CDF migration requirements.",
    )

    add_heading(doc, "OMS", "main")
    add_story(
        doc,
        "OMS001",
        "As a Data Engineer, I want all new CTT field values (Stakeholder Unit, Primary Technology, Funnel Stage, Technology/Campaign/Program, Offer Type) published to OMS tables, so that reporting teams can validate FY27 data elements from a centralized source.",
    )
    add_story(
        doc,
        "OMS002",
        "As a Reporting Analyst, I want OMS tables to reflect the updated CTT attribute model before reporting validation begins, so that sample reports use the same data source as future-state reporting.",
    )
    add_story(
        doc,
        "OMS003",
        "As a Data Engineer, I want the OMS tables to process upserts for modified CTT records, so that the data warehouse always reflects the latest edited state of an activation's attributes.",
    )

    add_heading(doc, "Summary of Impacts", "section")
    copy_table(source, 1, doc)

    add_heading(doc, "Appendix", "section")
    for idx, title in [
        (3, "Table 1 - ID Field List"),
        (4, "Table 2 - Stakeholder List"),
        (5, "Table 3 - Technology > Campaign > Program Mapping"),
        (6, "Table 4 - Funnel Stage"),
        (7, "Table 5 - Offer Category > Offer Type"),
        (8, "Table 6 - Business and Sub Business Entity"),
    ]:
        add_heading(doc, f"Appendix {idx - 2}", "main")
        add_para(doc, title, "Caption")
        copy_table(source, idx, doc)

    add_heading(doc, "Glossary of Terms", "section")
    add_heading(doc, "Glossary", "main")
    copy_table(source, 10, doc)

    add_heading(doc, "Approvals and Sign-Off", "section")
    add_heading(doc, "Document Reviewers", "main")
    copy_table(source, 11, doc)

    doc.save(BRD1_PATH)


def build_brd2(source):
    doc = blank_document_from_template(SOURCE_PATH)
    add_heading(doc, "Business Requirements Document", "section")
    build_metadata_block(
        doc,
        [
            ("Document Title", "Workfront to CTT Integration — Feasibility Study"),
            ("Document ID", "BRD-WF-CTT-002"),
            ("Version", "1.0 (Draft)"),
            ("Status", "Draft — For Review"),
            ("Target Delivery", "Q2 FY27 (earliest)"),
            ("Author", "Amanda Pattenden"),
            ("Date", "July 2026"),
            (
                "Related Documents",
                "20260729-Workfront to CTT Integration.docx; BRD-01 - CTT and CTT Request Portal FY27 Updates",
            ),
            (
                "Prerequisite",
                "BRD-01 CTT and CTT Request Portal FY27 updates must be deployed before integration feasibility validation",
            ),
        ],
    )
    build_revision_table(
        doc,
        "BRD-WF-CTT-002",
        "Workfront to CTT Integration Feasibility Study",
        "1.0",
        "Initial BRD split from combined integration document — Q2 feasibility study track",
    )

    add_heading(doc, "Introduction", "section")
    add_heading(doc, "Executive Summary", "main")
    add_para(
        doc,
        "This Business Requirements Document defines a feasibility study to assess whether Activity IDs (CCIDs) and Offer IDs (OIDs) can be automatically generated in the CTT tool directly from Adobe Workfront, with generated identifiers and FY27 attributes synchronized back into Workfront.",
    )
    add_para(
        doc,
        "The ideal end state is an interim integration that embeds ID creation into the Workfront workflow, removing the need for marketers to manually request tracking IDs through the CTT Request Portal or copy identifiers between systems.",
    )
    add_para(
        doc,
        "This work is targeted for Q2 at the earliest, pending completion of BRD-01 (CTT tool, CTT Request Portal, and OMS updates) and confirmation of technical prioritization and resourcing. This document reframes the prior combined POC scope as a structured feasibility study with clear decision criteria.",
    )

    add_heading(doc, "Problem Statement – Why?", "main")
    add_para(
        doc,
        "Manual ID administration creates adoption friction for Workfront teams. Marketers must understand CTT processes, submit requests through the CTT Request Portal, and maintain identifier consistency across systems—adding effort, error risk, and process complexity.",
    )
    add_para(
        doc,
        "Without automated integration, Workfront adoption may slow because teams cannot work entirely within Workfront without impacting campaign measurement continuity. An integration that triggers CTT ID creation from Workfront approvals would simplify workflows and reduce reliance on manual tracking processes.",
    )
    add_para(
        doc,
        "Before investing in build, the organization requires feasibility evidence covering technical connectivity, data mapping, write-back reliability, operational fit, and reporting continuity.",
    )

    add_heading(doc, "Vision – What?", "main")
    add_para(doc, "Primary Objective: Workfront & CTT Integration Feasibility Study")
    add_para(
        doc,
        "The feasibility study will validate whether Workfront-originated marketing activities can leverage CTT capabilities where legacy dependencies still exist. The target integration will:",
    )
    add_bullet(doc, "Automate Identifier Creation: Automatically generate Activity IDs (CCIDs) for channel activations and Offer IDs (OIDs) for content activations initiated through Workfront.")
    add_bullet(doc, "Synchronize Data Attributes: Synchronize generated identifiers and required FY27 reporting attributes between CTT and Workfront.")
    add_bullet(doc, "Ensure Consistent Data Governance: Ensure consistent capture of key reporting data elements across Workfront-managed activities.")
    add_bullet(doc, "Preserve Reporting Continuity: Maintain continuity for downstream reporting, attribution, and measurement processes that depend on CTT identifiers.")
    add_bullet(doc, "Simplify Marketing Workflows: Embed ID creation directly into the Workfront workflow to reduce manual CTT administration.")

    add_heading(doc, "Benefit Summary – Who?", "main")
    add_list(doc, "Marketing Teams: Automated CCID/OID creation removes manual ID request and copy steps.")
    add_list(doc, "Workfront Adoption Teams: Reduced friction for teams working primarily in Workfront.")
    add_list(doc, "Marketing Analytics & Insights Teams: Consistent FY27 attributes synchronized across systems.")
    add_list(doc, "Technology & Integration Teams: Validated integration approach and decommission evidence.")
    add_list(doc, "Program Leadership: Decision-quality evidence to scale, pause, or stop investment.")

    add_heading(doc, "Scope and Limitations", "section")
    add_heading(doc, "In-Scope", "main")
    add_bullet(doc, "Feasibility assessment of Workfront-triggered Activity ID and Offer ID creation in CTT.")
    add_bullet(doc, "Documentation of Workfront Task ID to Activity ID, Offer ID, Drive To ID, and UTM_ID relationships.")
    add_bullet(doc, "Pilot configuration of Workfront custom fields, Task ID references, and write-back attributes.")
    add_bullet(doc, "Bidirectional synchronization of agreed FY27 attributes between Workfront and CTT.")
    add_bullet(doc, "Integration error handling, field lock rules post-go-live, and discrepancy reporting.")
    add_bullet(doc, "Pilot activation type selection with known Workfront, CTT, and downstream dependencies.")
    add_bullet(doc, "Formal scale, pause, or stop recommendation with next-phase roadmap.")

    add_heading(doc, "Out of Scope", "main")
    add_list(doc, "CTT schema and CTT Request Portal updates (delivered under BRD-01).")
    add_list(doc, "OMS table schema changes beyond those delivered in BRD-01.")
    add_list(doc, "Full CTT decommissioning.")
    add_list(doc, "Salesforce, Eloqua, or OCC configuration changes.")
    add_list(doc, "Enterprise-wide Workfront rollout beyond the pilot use case.")
    add_list(doc, "Historical OCID data migration.")

    add_heading(doc, "Assumptions", "main")
    add_bullet(doc, "BRD-01 CTT and portal updates are deployed and stable before integration build begins.")
    add_bullet(doc, "Workfront Task IDs are available and usable as the primary reference identifier.")
    add_bullet(doc, "Endpoint connectivity exists or can be established between Workfront and CTT.")
    add_bullet(doc, "A defined pilot activation type with manageable dependencies is available.")
    add_bullet(doc, "Dedicated resources from Workfront, CTT/Integration, and Reporting teams will be assigned.")
    add_bullet(doc, "Integration remains a temporary bridge; CTT remains on the planned decommission path.")

    add_heading(doc, "Dependencies", "main")
    add_bullet(doc, "Completion and sign-off of BRD-01 deliverables.")
    add_bullet(doc, "Configuration of pilot custom fields and write-back trigger attributes in Workfront.")
    add_bullet(doc, "Availability of endpoint connectivity and mapping logic for automated ID generation.")
    add_bullet(doc, "Test environments, staging databases, and baseline datasets for validation.")
    add_bullet(doc, "Leadership alignment, funding confirmation, and Q2 prioritization decision.")

    add_heading(doc, "Success Metrics", "main")
    add_list(doc, "Integration Feasibility: Workfront can trigger CTT ID creation and receive generated IDs as usable data attributes.")
    add_list(doc, "Data Quality: ≥95% completeness for generated IDs and agreed FY27 attributes written back to Workfront.")
    add_list(doc, "User Experience: Pilot users can approve activations without additional manual tracking ID administration.")
    add_list(doc, "Synchronization Reliability: Field updates between Workfront and CTT reconcile within agreed tolerance.")
    add_list(doc, "Decision Outcome: Formal scale, pause, or stop recommendation supported by findings and risks.")

    add_heading(doc, "Stakeholder Requirements", "section")
    add_heading(doc, "Campaign Tagging & Tracking Tool", "main")
    add_heading(doc, "Automated ID Creation", "sub")
    add_story(
        doc,
        "CTT002",
        "As a CTT Administrator, I want Workfront-originated activation requests to automatically create Activity IDs for channel activations, so that the ID creation process is triggered without manual CTT Request Tool submission.",
    )
    add_story(
        doc,
        "CTT003",
        "As a CTT Administrator, I want Workfront-originated activation requests to automatically create Offer IDs for content activations, so that offer tracking remains consistent with current CTT operations.",
    )

    add_heading(doc, "Workfront", "main")
    for story in [
        (
            "WFT001",
            "As a Marketing Manager, I want Workfront to automatically trigger Activity ID creation in CTT when I approve a channel activation in the pilot scope using the data fields defined in Table 7 - Activity ID mapping Workfront - CTT, and have the generated Activity ID written back to the corresponding Workfront task.",
        ),
        (
            "WFT002",
            "As a Marketing Manager, I want Workfront to automatically trigger Offer ID creation in CTT when I approve a content activation in the pilot scope, so that consumed offers remain traceable without manual ID administration.",
        ),
        (
            "WFT003",
            "As a Workfront Administrator, I want generated Activity IDs and Offer IDs written back to Workfront as data attributes on the corresponding task/activation record, so that reporting teams can access IDs from Workfront data.",
        ),
        (
            "WFT004",
            "As a Workfront Administrator, I want field edits made to agreed FY27 reporting attributes on a Workfront record to automatically trigger an update to the corresponding record in CTT, so that both systems remain synchronized throughout the activation lifecycle.",
        ),
        (
            "WFT005",
            "As a Workfront Administrator, I want agreed FY27 attributes (Business Priority, Stakeholder Unit, Funnel Stage, campaign/program context, Task ID, UTM attributes) captured on Workfront records and available for integration.",
        ),
        (
            "WFT006",
            "As a CTT Administrator, I want Workfront to restrict or lock edits to synchronized fields once a channel or content activation goes Live (or after an ID has been pushed downstream), so that historical attribution data remains stable.",
        ),
        (
            "WFT007",
            "As an Integration Engineer, I want any failed field update synchronization between Workfront and CTT to trigger an automated alert to the Marketing Operations team, so that data discrepancies can be remediated before OMS ingestion.",
        ),
        (
            "WFT008",
            "As a Marketing Manager, I want to select and configure a pilot activation type with known dependencies, so that the feasibility study validates a realistic operating scenario without enterprise-wide rollout risk.",
        ),
    ]:
        add_story(doc, story[0], story[1])

    add_heading(doc, "Marketing Reporting / Data Validation", "main")
    add_story(
        doc,
        "RPT001",
        "As a Reporting Analyst, I want a weekly data discrepancy report to flag any Workfront Task IDs where data attributes in Workfront do not match the corresponding CTT record.",
    )
    add_story(
        doc,
        "RPT002",
        "As a Reporting Analyst, I want pilot reporting outputs validated against current attribution and funnel reporting expectations.",
    )

    add_heading(doc, "Enriched Requirements", "main")
    for story in [
        (
            "ER001",
            "As a Program Manager, I want a documented decommission blocker list with owner assignments, so that leadership has a credible path to CTT retirement beyond the pilot.",
        ),
        (
            "ER002",
            "As a Marketing Leader, I want a formal scale/pause/stop recommendation with decision criteria, so that investment beyond the feasibility study is evidence-based.",
        ),
        (
            "ER003",
            "As a Data Governance Lead, I want explicit exit criteria defining when CTT can be decommissioned, so that the bridge is not perceived as extending CTT life indefinitely.",
        ),
        (
            "ER004",
            "As a Workfront Adoption Lead, I want pilot user feedback on operational fit (manual steps, approval workflow, ID visibility), so that scale planning addresses adoption friction.",
        ),
        (
            "ER005",
            "As an Architecture Lead, I want the integration architecture assessed for scalability beyond the pilot, so that a scale decision does not require complete redesign.",
        ),
    ]:
        add_story(doc, story[0], story[1])

    add_heading(doc, "Feasibility Study Phases", "section")
    phases = [
        ("Phase 1 — Align", "Confirm pilot scope, success metrics, field mappings, and governance owners.", "Q2 Week 1-2"),
        ("Phase 2 — Assess", "Validate technical connectivity, API capabilities, and mapping completeness.", "Q2 Week 3-4"),
        ("Phase 3 — Build Pilot", "Configure Workfront triggers, CTT ID creation, and write-back for pilot activation type.", "Q2 Week 5-8"),
        ("Phase 4 — Validate", "Run pilot activations; reconcile data quality; prove reporting continuity.", "Q2 Week 9-10"),
        ("Phase 5 — Decide", "Deliver scale/pause/stop recommendation with next-phase backlog.", "Q2 Week 11-12"),
    ]
    table = doc.add_table(rows=1 + len(phases), cols=3)
    set_cell_text(table.rows[0].cells[0], "Phase")
    set_cell_text(table.rows[0].cells[1], "Activities")
    set_cell_text(table.rows[0].cells[2], "Indicative Timing")
    for i, (phase, activities, timing) in enumerate(phases, start=1):
        set_cell_text(table.rows[i].cells[0], phase)
        set_cell_text(table.rows[i].cells[1], activities)
        set_cell_text(table.rows[i].cells[2], timing)

    add_heading(doc, "Summary of Impacts", "section")
    impact_rows = [
        ("Workfront", "Configure pilot fields; trigger CTT ID creation; receive write-back of CCIDs/OIDs", "Integration"),
        ("CTT Activity ID", "Accept automated creation requests from Workfront for channel activations", "Integration"),
        ("CTT Offer ID", "Accept automated creation requests from Workfront for content activations", "Integration"),
        ("OMS Tables", "Ingest IDs and attributes created via Workfront-triggered flow", "Integration"),
        ("Marketing Reporting", "Validate synchronized Workfront/CTT pilot dataset and sample report", "Process"),
        ("Salesforce", "No changes during feasibility study", "No Change"),
        ("Eloqua", "No changes during feasibility study", "No Change"),
    ]
    table = doc.add_table(rows=1 + len(impact_rows), cols=3)
    set_cell_text(table.rows[0].cells[0], "Tool")
    set_cell_text(table.rows[0].cells[1], "Summary")
    set_cell_text(table.rows[0].cells[2], "CTT Direct Integration or Aligning Process")
    for i, row in enumerate(impact_rows, start=1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val)

    add_heading(doc, "Appendix", "section")
    add_heading(doc, "Appendix 1", "main")
    add_para(doc, "Table 7 - Activity ID mapping Workfront - CTT", "Caption")
    copy_table(source, 9, doc)

    add_heading(doc, "Glossary of Terms", "section")
    add_heading(doc, "Glossary", "main")
    glossary = [
        ("CCID", "Campaign Code ID — CTT Activity ID used for channel activations"),
        ("OID", "Offer ID — CTT identifier tracking content consumed by customers"),
        ("CTT", "Campaign Tagging and Tracking tool"),
        ("OMS", "Operational Marketing Store — centralized reporting data tables"),
        ("OCID", "Omni Channel ID — collective term for legacy identifiers (CCID, OID, DTID)"),
        ("Task ID", "Workfront task identifier"),
        ("UTM", "Urchin Tracking Module — standard web tracking parameters"),
    ]
    table = doc.add_table(rows=1 + len(glossary), cols=2)
    set_cell_text(table.rows[0].cells[0], "Acronym")
    set_cell_text(table.rows[0].cells[1], "Description")
    for i, (acr, desc) in enumerate(glossary, start=1):
        set_cell_text(table.rows[i].cells[0], acr)
        set_cell_text(table.rows[i].cells[1], desc)

    add_heading(doc, "Approvals and Sign-Off", "section")
    add_heading(doc, "Document Reviewers", "main")
    copy_table(source, 11, doc)

    doc.save(BRD2_PATH)


def main():
    source = Document(SOURCE_PATH)
    build_brd1(source)
    build_brd2(source)
    print(f"Created: {BRD1_PATH}")
    print(f"Created: {BRD2_PATH}")


if __name__ == "__main__":
    main()
