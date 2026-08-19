#!/usr/bin/env python3
"""Generate BRD v1.1 amendment section — Activity ID conditional intake logic."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = Path("/workspace/BRD-01-Amendment-v1.1-Conditional-Intake-Logic.docx")


def set_cell_shading(cell, fill="D9E2F3"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    return doc.add_paragraph(text, style=style)


def add_table(doc, headers, rows, header_fill="D9E2F3"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()
    return table


def add_story_block(doc, story_id, role, want, so_that, applies, priority, related):
    p = doc.add_paragraph()
    run = p.add_run(f"{story_id}")
    run.bold = True
    run.font.size = Pt(12)

    add_para(doc, f"As a {role},")
    add_para(doc, f"I want {want},")
    add_para(doc, f"so that {so_that}.")
    add_para(doc, f"Applies to: {applies}")
    add_para(doc, f"Priority: {priority}")
    add_para(doc, f"Related: {related}")
    doc.add_paragraph()


def add_ac_heading(doc, ac_id, title):
    p = doc.add_paragraph()
    run = p.add_run(f"{ac_id} — {title}")
    run.bold = True


def build_document():
    doc = Document()

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Business Requirements Document — Amendment")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Activity ID Conditional Intake Logic (v1.1)")
    run.font.size = Pt(14)

    doc.add_paragraph()

    meta = [
        ("Parent Document", "BRD-01 — CTT Tool and CTT Request Portal FY27 Updates"),
        ("Document ID", "BRD-WF-CTT-001-AMD-1.1"),
        ("Version", "1.1 (Amendment)"),
        ("Status", "Draft — For Review"),
        ("Author", "Amanda Pattenden"),
        ("Date", date.today().strftime("%d %B %Y")),
        (
            "Related Documents",
            "BRD-01 v1.0; Request Portal Business Logic — Business rules (conditional intake logic)",
        ),
        ("Purpose", "Insert or append this section to BRD-01 as the v1.1 controlled amendment."),
    ]
    add_table(doc, ["Field", "Value"], meta, header_fill="BDD7EE")

    # Revision History
    add_heading(doc, "Revision History", 1)
    add_table(
        doc,
        ["Date", "Version", "Description", "Author"],
        [
            (
                date.today().strftime("%d/%m/%Y"),
                "1.1",
                "Amendment: Activity ID conditional intake logic (Campaign-first; max two user "
                "selections; mandatory auto-population of third field). Supersedes v1.0 selection "
                "order in POR002. Source: Request Portal Business Logic — Business rules "
                "(conditional intake logic).",
                "Amanda Pattenden",
            )
        ],
    )

    # Change Summary
    add_heading(doc, "Document Change Summary — v1.1", 1)
    add_para(
        doc,
        "The following changes amend BRD-01 v1.0. Requirements marked Superseded must not be implemented.",
    )
    add_table(
        doc,
        ["Change ID", "Type", "Requirement", "Summary of Change", "Business Driver"],
        [
            (
                "CR-01",
                "Amend",
                "POR002",
                "Replace fixed Technology > Campaign > Program selection order with Campaign-first "
                "conditional intake. Requester answers at most two questions; the third field always "
                "auto-populates.",
                "Reduce requester effort; enforce valid Campaign / Program / Primary Technology "
                "combinations for FY27 reporting",
            ),
            (
                "CR-02",
                "New",
                "POR010",
                "Mandate that Campaign, Program, and Primary Technology are always populated on "
                "submit; blanks are data defects.",
                "FY27 reporting integrity; no incomplete hierarchy records in CTT or OMS",
            ),
        ],
    )

    add_heading(doc, "Superseded Text (v1.0 — Do Not Implement)", 2)
    p = doc.add_paragraph()
    run = p.add_run(
        "POR002 (v1.0): “Selection hierarchy changed from Buying Centre > Campaign Theme > Program "
        "to Technology > Campaign > Program across all Activity ID request types.”"
    )
    run.italic = True

    # Scope amendments
    add_heading(doc, "Scope Amendments", 1)
    add_heading(doc, "In-Scope — Replace Existing Bullet", 2)
    add_para(
        doc,
        "Replace: “Update Activity ID schema: … Technology > Campaign > Program hierarchy.”",
        italic=True,
    )
    add_para(
        doc,
        "With: “Update Activity ID schema and CTT Request Portal conditional intake logic: "
        "Campaign-first presentation; dependent Primary Technology or Program selection (max two "
        "user inputs); automatic population of the remaining field(s) per approved campaign matrix.”",
        bold=True,
    )

    add_heading(doc, "In-Scope — Add", 2)
    add_bullet(
        doc,
        "Conditional intake business rules for Activity ID Create New workflow, including filtered "
        "Program picklists and campaign-specific auto-population of Primary Technology and/or Program.",
    )

    add_heading(doc, "Assumptions — Add", 2)
    add_bullet(
        doc,
        "The approved Campaign → Program → Primary Technology matrix (Table 3 and conditional intake "
        "rules) is stable for Q1 configuration. Changes to the matrix require a further BRD amendment.",
    )

    # POR002
    add_heading(doc, "Amended Requirement — POR002", 1)
    add_story_block(
        doc,
        "POR002 (Amended v1.1) — Activity ID Conditional Intake Hierarchy",
        "Marketing Manager",
        "the CTT Request Portal Activity ID request flow to present Campaign first, then show only "
        "the next field required based on my Campaign selection (Primary Technology or Program), with "
        "the remaining field(s) auto-populated",
        "I capture a valid FY27 Campaign / Program / Primary Technology combination with minimal input "
        "and without invalid or incomplete hierarchy data",
        "Create New Activity ID (primary). Create Bulk and Update Existing: see acceptance criteria below.",
        "Must Have",
        "POR010, Table 3 (Technology > Campaign > Program Mapping — matrix maintained as reference data)",
    )

    add_heading(doc, "Acceptance Criteria", 2)

    add_ac_heading(doc, "AC-POR002-01", "Campaign is Always Step 1")
    add_bullet(doc, "Campaign is always the first field presented to the requester.")
    add_bullet(doc, "Campaign is required before any dependent field is shown or auto-populated.")

    add_ac_heading(doc, "AC-POR002-02", "Maximum Two User Questions")
    add_bullet(
        doc,
        "The requester answers a maximum of two questions: (1) Campaign, and (2) either Primary "
        "Technology or Program, depending on Campaign.",
    )
    add_bullet(
        doc,
        "The third hierarchy field always auto-populates; the requester does not manually enter it "
        "when intake logic applies.",
    )

    add_ac_heading(doc, "AC-POR002-03", "Path A — Primary Technology Picker")
    add_para(
        doc,
        "When Campaign is No Campaign Aligned or Product Launch:",
    )
    add_table(
        doc,
        ["User Action", "System Behaviour"],
        [
            ("User selects Campaign", "Show Primary Technology picker (Step 2)"),
            ("User selects Primary Technology", "Auto-set Program = No Program Aligned"),
            ("Second question shown?", "Yes — Primary Technology"),
        ],
    )

    add_ac_heading(doc, "AC-POR002-04", "Path B — Filtered Program Picker")
    add_para(
        doc,
        "When Campaign is: Up Market - Modernize and Secure Network; Mid Market - Modernize and "
        "Secure Network; DC Modernization; or Full-Stack AI Infrastructure:",
    )
    add_table(
        doc,
        ["User Action", "System Behaviour"],
        [
            ("User selects Campaign", "Show Program picker with filtered list only (see AC-POR002-06)"),
            ("User selects Program", "Auto-set Primary Technology per matrix below"),
            ("Second question shown?", "Yes — Program"),
        ],
    )
    add_para(doc, "Primary Technology auto-population (Path B):", bold=True)
    add_table(
        doc,
        ["Campaign", "Primary Technology (auto)"],
        [
            ("Up Market - Modernize and Secure Network", "Networking"),
            ("Mid Market - Modernize and Secure Network", "Networking"),
            ("DC Modernization", "Data Center"),
            ("Full-Stack AI Infrastructure", "Data Center"),
        ],
    )

    add_ac_heading(doc, "AC-POR002-05", "Path C — No Second Question")
    add_para(
        doc,
        "When Campaign is: SMB; Partner; CxO; Secure Networking; AI Brand Campaign; Project Beacon; "
        "Sovereign; or Neocloud:",
    )
    add_table(
        doc,
        ["User Action", "System Behaviour"],
        [
            ("User selects Campaign", "No further picker — intake complete"),
            ("Program (auto)", "No Program Aligned"),
            ("Primary Technology (auto)", "Per matrix below"),
            ("Second question shown?", "No"),
        ],
    )
    add_para(doc, "Auto-population (Path C):", bold=True)
    add_table(
        doc,
        ["Campaign", "Program (auto)", "Primary Technology (auto)"],
        [
            ("SMB", "No Program Aligned", "Networking"),
            ("Partner", "No Program Aligned", "Cross Technology"),
            ("CxO", "No Program Aligned", "Cross Technology"),
            ("Secure Networking", "No Program Aligned", "Cross Technology"),
            ("AI Brand Campaign", "No Program Aligned", "Cross Technology"),
            ("Project Beacon", "No Program Aligned", "Cross Technology"),
            ("Sovereign", "No Program Aligned", "Cross Technology"),
            ("Neocloud", "No Program Aligned", "Service Provider"),
        ],
    )

    add_ac_heading(doc, "AC-POR002-06", "Filtered Program Lists (Path B Only)")
    add_para(doc, "When the Program picker is shown, only the following options are available:")
    add_table(
        doc,
        ["Campaign", "Program Options"],
        [
            (
                "Up Market - Modernize and Secure Network",
                "No Program Aligned; Up Market - Wireless LDOS; Up Market - SASE Attach; "
                "Up Market - Industry Vertical; Up Market - Competitor",
            ),
            (
                "Mid Market - Modernize and Secure Network",
                "No Program Aligned; Mid Market - Less uh-ohs. More ah-has.; "
                "Mid Market - Industry Vertical; Mid Market - Competitor",
            ),
            (
                "DC Modernization",
                "No Program Aligned; DC Modernization - Industry Vertical",
            ),
            (
                "Full-Stack AI Infrastructure",
                "No Program Aligned; NVIDIA; Full-Stack AI Infrastructure - Industry Vertical; "
                "Full-Stack AI Infrastructure - Competitor",
            ),
        ],
    )
    add_bullet(
        doc,
        "Programs not in the filtered list for the selected Campaign must not be selectable in Create New intake.",
    )

    add_ac_heading(doc, "AC-POR002-07", "Field Visibility")
    add_bullet(
        doc,
        "Auto-populated fields (Program and/or Primary Technology) must be visible to the requester "
        "after Campaign selection (and Step 2 where applicable), showing the system-assigned values "
        "before submit.",
    )
    add_bullet(
        doc,
        "Auto-populated values must be editable only if explicitly allowed in a future BRD amendment; "
        "v1.1 default: not editable after auto-population (confirm with business if override is required).",
    )

    add_ac_heading(doc, "AC-POR002-08", "Create Bulk and Update Existing")
    add_bullet(
        doc,
        "Create Bulk: Bulk template and validation must enforce the same Campaign / Program / Primary "
        "Technology combinations; invalid combinations must fail validation with a clear error. "
        "Auto-population rules may be applied at row validation time per the same matrix.",
    )
    add_bullet(
        doc,
        "Update Existing: Existing records may retain legacy combinations until edited; on save, all "
        "three fields must be populated (POR010). If the user changes Campaign, re-apply conditional "
        "intake logic (reset dependent fields per Path A/B/C).",
    )
    add_para(
        doc,
        "Traceability: Request Portal Business Logic — Business rules (conditional intake logic)",
        italic=True,
    )

    # POR010
    add_heading(doc, "New Requirement — POR010", 1)
    add_story_block(
        doc,
        "POR010 (New v1.1) — Mandatory Population of Hierarchy Fields",
        "Data Governance Lead",
        "Campaign, Program, and Primary Technology to always be populated on every Activity ID request at submission",
        "FY27 reporting and OMS-published data never contain blank hierarchy attributes (data defects)",
        "Create New, Create Bulk, Update Existing — Activity ID request types",
        "Must Have",
        "POR002, OMS001",
    )

    add_heading(doc, "Acceptance Criteria", 2)
    add_ac_heading(doc, "AC-POR010-01", "No Blanks on Submit")
    add_bullet(
        doc,
        "On submit (Create New), row validation (Create Bulk), or save (Update Existing), the system "
        "must reject any Activity ID where Campaign, Program, or Primary Technology is null, empty, or unset.",
    )
    add_ac_heading(doc, "AC-POR010-02", "Data Defect Definition")
    add_bullet(
        doc,
        "Any Activity ID persisted or published to OMS with a missing Campaign, Program, or Primary "
        "Technology is a data defect and must be flagged for remediation (reporting / admin queue — "
        "implementation detail for FRD).",
    )
    add_ac_heading(doc, "AC-POR010-03", "Auto-Population Satisfies Mandatory Rule")
    add_bullet(
        doc,
        "Values set by conditional intake auto-population (Path A, B, C) count as populated for "
        "submission and OMS publish.",
    )
    add_para(
        doc,
        "Traceability: Request Portal Business Logic — Key requirement: All three fields must always "
        "be populated. Blanks are data defects.",
        italic=True,
    )

    # Decision matrix
    add_heading(doc, "Appendix A — Activity ID Intake Decision Matrix (Authoritative for v1.1)", 1)
    add_table(
        doc,
        ["#", "Campaign", "Step 2 (User)", "Program (Result)", "Primary Technology (Result)", "Path"],
        [
            ("1", "No Campaign Aligned", "Primary Technology (picker)", "No Program Aligned (auto)", "User selection", "A"),
            ("2", "Product Launch", "Primary Technology (picker)", "No Program Aligned (auto)", "User selection", "A"),
            ("3", "Up Market - Modernize and Secure Network", "Program (filtered)", "User selection", "Networking (auto)", "B"),
            ("4", "Mid Market - Modernize and Secure Network", "Program (filtered)", "User selection", "Networking (auto)", "B"),
            ("5", "DC Modernization", "Program (filtered)", "User selection", "Data Center (auto)", "B"),
            ("6", "Full-Stack AI Infrastructure", "Program (filtered)", "User selection", "Data Center (auto)", "B"),
            ("7", "SMB", "—", "No Program Aligned (auto)", "Networking (auto)", "C"),
            ("8", "Partner", "—", "No Program Aligned (auto)", "Cross Technology (auto)", "C"),
            ("9", "CxO", "—", "No Program Aligned (auto)", "Cross Technology (auto)", "C"),
            ("10", "Secure Networking", "—", "No Program Aligned (auto)", "Cross Technology (auto)", "C"),
            ("11", "AI Brand Campaign", "—", "No Program Aligned (auto)", "Cross Technology (auto)", "C"),
            ("12", "Project Beacon", "—", "No Program Aligned (auto)", "Cross Technology (auto)", "C"),
            ("13", "Sovereign", "—", "No Program Aligned (auto)", "Cross Technology (auto)", "C"),
            ("14", "Neocloud", "—", "No Program Aligned (auto)", "Service Provider (auto)", "C"),
        ],
    )

    add_heading(doc, "Three Paths — Summary for Stakeholders", 2)
    add_table(
        doc,
        ["Path", "Trigger", "User Answers", "Auto-Populated"],
        [
            ("A", "No Campaign Aligned, Product Launch", "Campaign + Primary Technology", "Program → No Program Aligned"),
            ("B", "Campaigns with multiple Program options", "Campaign + Program (filtered)", "Primary Technology (from matrix)"),
            ("C", "Campaigns with no second question", "Campaign only", "Program + Primary Technology (both from matrix)"),
        ],
    )

    # UAT
    add_heading(doc, "Appendix B — UAT Scenarios", 1)
    add_table(
        doc,
        ["ID", "Scenario", "Steps", "Expected Result"],
        [
            ("UAT-POR002-01", "Path A", "Select Product Launch → select Primary Technology Networking", "Program = No Program Aligned; submit succeeds"),
            ("UAT-POR002-02", "Path B", "Select Up Market - Modernize and Secure Network → select Up Market - SASE Attach", "Primary Technology = Networking; only filtered Programs shown"),
            ("UAT-POR002-03", "Path B filter", "Select DC Modernization", "Program list contains only: No Program Aligned, DC Modernization - Industry Vertical"),
            ("UAT-POR002-04", "Path C", "Select SMB", "No Step 2; Program = No Program Aligned; Primary Technology = Networking"),
            ("UAT-POR002-05", "Path C Neocloud", "Select Neocloud", "Program = No Program Aligned; Primary Technology = Service Provider"),
            ("UAT-POR010-01", "Data defect", "Attempt submit with Campaign set but Program missing", "Submit blocked; error indicates mandatory field"),
            ("UAT-POR002-06", "Campaign change", "Path B selection → change Campaign to SMB", "Step 2 hidden; Program and Primary Technology reset per Path C"),
        ],
    )

    # Cross-reference
    add_heading(doc, "Appendix C — Cross-Reference and Open Points", 1)
    add_heading(doc, "CTT007 Cross-Reference", 2)
    add_para(
        doc,
        "Amended v1.1: Request Portal Activity ID intake is Campaign-first (POR002). CTT007 applies to "
        "CTT tool Activity ID UI/API behaviour unless explicitly aligned to Portal rules in FRD.",
    )

    add_heading(doc, "Open Points for Business Confirmation", 2)
    add_bullet(doc, "Editable auto-populated fields — AC-POR002-07 assumes read-only after auto-fill; confirm if marketers may override.")
    add_bullet(
        doc,
        "Path C wording — “only one program” in the business document maps to campaigns with no Step 2 "
        "(fixed Program + Primary Technology); confirm naming with stakeholders.",
    )
    add_bullet(
        doc,
        "Offer ID (POR008) — unchanged by this amendment; confirm if Offer ID should mirror Campaign-first logic in a future amendment.",
    )

    add_heading(doc, "Note to CTT Request Portal Team", 2)
    add_para(
        doc,
        "BRD v1.1 amends POR002 and adds POR010. v1.0 Technology-first hierarchy is withdrawn for "
        "Activity ID Create New. Detailed rules are in the acceptance criteria above; source document: "
        "Request Portal Business Logic — conditional intake. Please confirm bulk template and Update "
        "Existing re-validation approach before build.",
    )

    doc.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
